import pandas as pd
import tkinter as tk
from tkinter import messagebox
from tkinter import ttk

#Instalar extensão 'docx' e 'python-docx'
from docx import Document
from docx.shared import Pt



janela = tk.Tk()
janela.title("Gerador de certificados")

estilo = ttk.Style()
estilo.theme_use("alt")
estilo.configure(".", font="Arial 20", rowheight=30)

treeviewDados = ttk.Treeview(janela, columns=(1, 2, 3, 4, 5, 6), show="headings")

treeviewDados.column("1", anchor="center")
treeviewDados.heading("1", text="CPF")

treeviewDados.column("2", anchor="center")
treeviewDados.heading("2", text="Nome")

treeviewDados.column("3", anchor="center")
treeviewDados.heading("3", text="RG")

treeviewDados.column("4", anchor="center")
treeviewDados.heading("4", text="Data Inicio")

treeviewDados.column("5", anchor="center")
treeviewDados.heading("5", text="Data Fim")

treeviewDados.column("6", anchor="center")
treeviewDados.heading("6", text="Email")

treeviewDados.grid(row=4, column=0, columnspan=6, sticky="NSEW", pady=15)

def funcaoPassaDadosTreeviewParaEntry(event):

    #Pego o item selecionado
    item = treeviewDados.selection()

    for i in item:

        #Limpando os campos de entrada de dados
        exibirCPF.delete(0, "end")
        exibirNome.delete(0, "end")
        exibirRG.delete(0, "end")
        exibirDataInicio.delete(0, "end")
        exibirDataFim.delete(0, "end")
        exibirEmail.delete(0, "end")

        exibirCPF.insert(0, treeviewDados.item(i, "values")[0])
        exibirNome.insert(0, treeviewDados.item(i, "values")[1])
        exibirRG.insert(0, treeviewDados.item(i, "values")[2])
        exibirDataInicio.insert(0, treeviewDados.item(i, "values")[3])
        exibirDataFim.insert(0, treeviewDados.item(i, "values")[4])
        exibirEmail.insert(0, treeviewDados.item(i, "values")[5])

#Passa os dados da linha da treeview clicada 2x para os campos Entry
treeviewDados.bind("<Double-1>", funcaoPassaDadosTreeviewParaEntry)


dadosUsuarios = pd.read_excel("C:\\Users\\roti5\\OneDrive\\Documentos\\PROGRAMAÇÃO\\Curso lógica de programação - Udemy\\TKinter\\Certificados\\Dados.xlsx")

#Convertendo a coluna Data Inicio para o tipo texto (str)
dadosUsuarios["Data Inicio"] = dadosUsuarios["Data Inicio"].astype(str)

#Convertendo a coluna Data Fim para o tipo texto (str)
dadosUsuarios["Data Fim"] = dadosUsuarios["Data Fim"].astype(str)



#iloc - permite pegar o texto
for linha in range(len(dadosUsuarios)):

    #Tratando o formato da DataInicio
    dataInicioAno = dadosUsuarios.iloc[linha, 4].split("-")[0] #.iloc - pegando o texto da coluna 3 que é a data // .split - quebra a coluna em partes de acordo com um critério '-' // [0] - pega a primeira parte
    dataInicioMes = dadosUsuarios.iloc[linha, 4].split("-")[1]
    dataInicioDia = dadosUsuarios.iloc[linha, 4].split("-")[2]

    dataInicioTratada = dataInicioDia + "/" + dataInicioMes + "/" + dataInicioAno

    # Tratando o formato da DataFim
    dataFimAno = dadosUsuarios.iloc[linha, 3].split("-")[0]
    dataFimMes = dadosUsuarios.iloc[linha, 3].split("-")[1]
    dataFimDia = dadosUsuarios.iloc[linha, 3].split("-")[2]

    dataFimTratada = dataFimDia + "/" + dataFimMes + "/" + dataFimAno

    #Populando a Treeview com os dados do Excel
    treeviewDados.insert("", "end",
                         values=(str(dadosUsuarios.iloc[linha, 0]), #iloc - permite pegar o texto da coluna inserida (0=CPF)
                                 str(dadosUsuarios.iloc[linha, 1]), #Pega o nome
                                 str(dadosUsuarios.iloc[linha, 2]), #Pega o RG
                                 str(dataInicioTratada), #Pega Data Inicio
                                 str(dataFimTratada), #Pega Data Fim
                                 str(dadosUsuarios.iloc[linha, 5]))) #Pega o Email

cpf = tk.Label(text="CPF: ", font="Arial 12")
cpf.grid(row=0, column=0, sticky="E", pady=15)
exibirCPF = tk.Entry(font="Arial 12")
exibirCPF.grid(row=0, column=1, sticky="W", pady=15)

nome = tk.Label(text="Nome: ", font="Arial 12")
nome.grid(row=0, column=2, sticky="E", pady=15)
exibirNome = tk.Entry(font="Arial 12")
exibirNome.grid(row=0, column=3, sticky="W", pady=15)

rg = tk.Label(text="RG: ", font="Arial 12")
rg.grid(row=0, column=4, sticky="E", pady=15)
exibirRG = tk.Entry(font="Arial 12")
exibirRG.grid(row=0, column=5, sticky="W", pady=15)

dataInicio = tk.Label(text="Data Inicio: ", font="Arial 12")
dataInicio.grid(row=1, column=0, sticky="E", pady=15)
exibirDataInicio = tk.Entry(font="Arial 12")
exibirDataInicio.grid(row=1, column=1, sticky="W", pady=15)

dataFim = tk.Label(text="Data Fim: ", font="Arial 12")
dataFim.grid(row=1, column=2, sticky="E", pady=15)
exibirDataFim = tk.Entry(font="Arial 12")
exibirDataFim.grid(row=1, column=3, sticky="W", pady=15)

email = tk.Label(text="Email: ", font="Arial 12")
email.grid(row=1, column=4, sticky="E", pady=15)
exibirEmail = tk.Entry(font="Arial 12")
exibirEmail.grid(row=1, column=5, sticky="W", pady=15)

def filtrarDados():

    for linha in range(len(dadosUsuarios)):

        todasLinhas = treeviewDados.get_children()

        #Deletando todas as linhas da treeview
        treeviewDados.delete(*todasLinhas)

        #Se o 'CPF for vazio', não tiver nenhuma informação eu carrego todos os dados na treeview
        if exibirCPF.get() == "":

            botaoPesquisar.config(text="Filtrar") #Mudando o texto do botão

            for linha in range(len(dadosUsuarios)):

                # Tratando o formato da DataInicio
                dataInicioAno = dadosUsuarios.iloc[linha, 4].split("-")[0]  # .iloc - pegando o texto da coluna 3 que é a data // .split - quebra a coluna em partes de acordo com um critério '-' // [0] - pega a primeira parte
                dataInicioMes = dadosUsuarios.iloc[linha, 4].split("-")[1]
                dataInicioDia = dadosUsuarios.iloc[linha, 4].split("-")[2]

                dataInicioTratada = dataInicioDia + "/" + dataInicioMes + "/" + dataInicioAno

                # Tratando o formato da DataFim
                dataFimAno = dadosUsuarios.iloc[linha, 3].split("-")[0]
                dataFimMes = dadosUsuarios.iloc[linha, 3].split("-")[1]
                dataFimDia = dadosUsuarios.iloc[linha, 3].split("-")[2]

                dataFimTratada = dataFimDia + "/" + dataFimMes + "/" + dataFimAno

                # Populando a Treeview com os dados do Excel
                treeviewDados.insert("", "end",
                                     values=(str(dadosUsuarios.iloc[linha, 0]),
                                             # iloc - permite pegar o texto da coluna inserida (0=CPF)
                                             str(dadosUsuarios.iloc[linha, 1]),  # Pega o nome
                                             str(dadosUsuarios.iloc[linha, 2]),  # Pega o RG
                                             str(dataInicioTratada),  # Pega Data Inicio
                                             str(dataFimTratada),  # Pega Data Fim
                                             str(dadosUsuarios.iloc[linha, 5])))  # Pega o Email

                exibirCPF.delete(0, "end")
                exibirNome.delete(0, "end")
                exibirRG.delete(0, "end")
                exibirDataInicio.delete(0, "end")
                exibirDataFim.delete(0, "end")
                exibirEmail.delete(0, "end")

        else:

            botaoPesquisar.config(text="Limpar Filtros")  # Mudando o texto do botão

            for linha in range(len(dadosUsuarios)):

                #Verifico se o CPF do campo Entry é  igual ao CPF da linha corrente do treeview
                if exibirCPF.get() == str(dadosUsuarios.iloc[linha, 0]):

                    # Tratando o formato da DataInicio
                    dataInicioAno = dadosUsuarios.iloc[linha, 4].split("-")[
                        0]  # .iloc - pegando o texto da coluna 3 que é a data // .split - quebra a coluna em partes de acordo com um critério '-' // [0] - pega a primeira parte
                    dataInicioMes = dadosUsuarios.iloc[linha, 4].split("-")[1]
                    dataInicioDia = dadosUsuarios.iloc[linha, 4].split("-")[2]

                    dataInicioTratada = dataInicioDia + "/" + dataInicioMes + "/" + dataInicioAno

                    # Tratando o formato da DataFim
                    dataFimAno = dadosUsuarios.iloc[linha, 3].split("-")[0]
                    dataFimMes = dadosUsuarios.iloc[linha, 3].split("-")[1]
                    dataFimDia = dadosUsuarios.iloc[linha, 3].split("-")[2]

                    dataFimTratada = dataFimDia + "/" + dataFimMes + "/" + dataFimAno

                    # Populando a Treeview com os dados do Excel
                    treeviewDados.insert("", "end",
                                         values=(str(dadosUsuarios.iloc[linha, 0]),
                                                 # iloc - permite pegar o texto da coluna inserida (0=CPF)
                                                 str(dadosUsuarios.iloc[linha, 1]),  # Pega o nome
                                                 str(dadosUsuarios.iloc[linha, 2]),  # Pega o RG
                                                 str(dataInicioTratada),  # Pega Data Inicio
                                                 str(dataFimTratada),  # Pega Data Fim
                                                 str(dadosUsuarios.iloc[linha, 5])))  # Pega o Email







botaoPesquisar = tk.Button(text="Pesquisar", font="Arial 14", command=filtrarDados)
botaoPesquisar.grid(row=5, column=0, columnspan=2, sticky="NSEW", padx=20)

def gerarCertificado():

    arquivoWord = Document("Certificado.docx")
    estilo = arquivoWord.styles["Normal"]

    #Pegando os dados do aluno dos campos Entry
    nomeAluno = exibirNome.get()
    dataInicio = exibirDataInicio.get()
    dataFim = exibirDataFim.get()
    nomeInstrutor = "Clevison Santos"
    cpf_Aluno = exibirCPF.get()
    rg_Aluno = exibirRG.get()

    frase_parte1 = " concluiu com sucesso o curso de Python RPA, com a carga horária de 20 horas, promovido pela escola de Cursos Online de "
    frase_montada = f" cujo CPF é {cpf_Aluno} e RG de número: {rg_Aluno},{frase_parte1} {dataInicio} a {dataFim}."

    for paragrafo in arquivoWord.paragraphs:

        if "@nome" in paragrafo.text:
            paragrafo.text = f"{nomeAluno}," #Substitui todo o parágrafo pelo texto criado na variável
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)

        if "@DataFim" in paragrafo.text:
            paragrafo.text = frase_montada  #Substitui todo o parágrafo pelo texto criado na variável
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)

    #Montando o caminho + nome do certificado que será criado para cada aluno
    caminhoCertificadoGerado = "C:\\Users\\roti5\\OneDrive\\Documentos\\PROGRAMAÇÃO\\Curso lógica de programação - Udemy\\TKinter\\Certificados\\" + nomeAluno + ".docx"

    #Salvando o certificado como o nome do aluno
    arquivoWord.save(caminhoCertificadoGerado)

    # Limpando os campos de entrada de dados da pessoa que teve o certificado gerado
    exibirCPF.delete(0, "end")
    exibirNome.delete(0, "end")
    exibirRG.delete(0, "end")
    exibirDataInicio.delete(0, "end")
    exibirDataFim.delete(0, "end")
    exibirEmail.delete(0, "end")

    messagebox.showinfo("Mensagem", "Certificado gerado com sucesso!")

#Gera o certificado das informações que estão no campo Entry (Clicar 2x na pessoa...)
botaoGerarCertificado = tk.Button(text="Gerar Certificado", font="Arial 14", command=gerarCertificado)
botaoGerarCertificado.grid(row=5, column=2, columnspan=2, sticky="NSEW", padx=20)

def gerarCertificadoMassa():

    for linha in treeviewDados.get_children(): #Passa linha por linha da treeview

        coluna = treeviewDados.item(linha)["values"] #Passa os valores da linha corrente e transforma em colunas

        CPF_Separado = coluna[0] #Pega as informações de cada coluna criada na linha anterior
        nomeAluno_Separado = coluna[1]
        RG_Separado = coluna[2]
        dataInicio_Separado = coluna[3]
        dataFim_Separado = coluna[4]
        nomeInstrutor_Separado = "Clevison Santos"

        arquivoWord = Document("Certificado.docx")
        estilo = arquivoWord.styles["Normal"]


        frase_parte1 = " concluiu com sucesso o curso de Python RPA, com a carga horária de 20 horas, promovido pela escola de Cursos Online de"
        frase_montada = f" cujo CPF é {CPF_Separado} e RG de número: {RG_Separado},{frase_parte1} {dataInicio_Separado} a {dataFim_Separado}."

        for paragrafo in arquivoWord.paragraphs:

            if "@nome" in paragrafo.text:
                paragrafo.text = f"{nomeAluno_Separado},"  # Substitui todo o parágrafo pelo texto criado na variável
                fonte = estilo.font
                fonte.name = "Calibri (Corpo)"
                fonte.size = Pt(24)

            if "@DataFim" in paragrafo.text:
                paragrafo.text = frase_montada  # Substitui todo o parágrafo pelo texto criado na variável
                fonte = estilo.font
                fonte.name = "Calibri (Corpo)"
                fonte.size = Pt(24)

        # Montando o caminho + nome do certificado que será criado para cada aluno
        caminhoCertificadoGerado = "C:\\Users\\roti5\\OneDrive\\Documentos\\PROGRAMAÇÃO\\Curso lógica de programação - Udemy\\TKinter\\Certificados\\" + nomeAluno_Separado + ".docx"

        # Salvando o certificado como o nome do aluno
        arquivoWord.save(caminhoCertificadoGerado)

    messagebox.showinfo("Mensagem", "Certificados gerados com sucesso!")



botaoGerarCertificadoEmMassa = tk.Button(text="Gerar em Massa", font="Arial 14", command=gerarCertificadoMassa)
botaoGerarCertificadoEmMassa.grid(row=5, column=4, columnspan=2, sticky="NSEW", padx=20)

janela.mainloop()











